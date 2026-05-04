"""File-attachment handler.

The node is invoked when the router detects a document attachment on the
incoming request. We dispatch on file extension to the right parser
(``pdfplumber`` / ``python-docx`` / ``openpyxl``), normalise the result
into a :class:`ParsedFile`, then ask a light model to summarise the
document with a focus on data tables and emissions-related figures.

The parser layer is intentionally separated from the node layer so that
the helpers can be unit-tested with raw paths and re-used by future
batch-ingestion flows.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Literal, cast

from langchain_core.messages import AIMessage
from langchain_core.runnables import RunnableConfig
from pydantic import BaseModel, Field

from app.models.state import SharedState as AgentState
from app.utils.llm import get_light_model, safe_ainvoke, strip_thinking

_log = logging.getLogger(__name__)

FileKind = Literal["pdf", "docx", "xlsx"]

# Caps to keep memory + LLM token usage bounded.
_MAX_TEXT_CHARS = 24_000  # text passed to summariser
_MAX_PDF_PAGES = 30
_MAX_DOCX_PARAGRAPHS = 1500
_MAX_XLSX_ROWS_PER_SHEET = 200
_FALLBACK_REPLY = (
    "文件已上传，但解析失败或暂不支持该格式。"
    "当前支持 .pdf / .docx / .xlsx，请确认后重试。"
)


class ParsedFile(BaseModel):
    """Normalised representation of a parsed document.

    Attributes:
        kind: Source format. One of ``pdf``, ``docx``, ``xlsx``.
        pages: Logical page count (PDF), paragraph count (DOCX), or sheet
            count (XLSX). Used for diagnostic logging.
        text: Concatenated free text. Always trimmed to
            :data:`_MAX_TEXT_CHARS` before being handed to the LLM.
        tables: Tables as 3-D lists ``[table][row][cell]``. Empty for
            documents that contain no tabular content.
    """

    kind: FileKind
    pages: int
    text: str
    tables: list[list[list[str]]] = Field(default_factory=list)


# --------------------------------------------------------------------------- #
# Parsers                                                                     #
# --------------------------------------------------------------------------- #


def _parse_pdf(path: Path) -> ParsedFile:
    """Parse a PDF using ``pdfplumber``.

    Pages beyond :data:`_MAX_PDF_PAGES` are dropped to keep latency bounded.
    Both running text and embedded tables are extracted.
    """
    import pdfplumber

    text_parts: list[str] = []
    tables: list[list[list[str]]] = []
    page_count = 0

    with pdfplumber.open(str(path)) as pdf:
        for idx, page in enumerate(pdf.pages):
            if idx >= _MAX_PDF_PAGES:
                text_parts.append("... (后续页面已截断)")
                break
            page_count = idx + 1
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[第 {idx + 1} 页]\n{page_text.strip()}")
            for raw_table in page.extract_tables() or []:
                normalised = [
                    ["" if cell is None else str(cell).strip() for cell in row]
                    for row in raw_table
                ]
                if any(any(c for c in row) for row in normalised):
                    tables.append(normalised)

    return ParsedFile(
        kind="pdf",
        pages=page_count,
        text="\n\n".join(text_parts),
        tables=tables,
    )


def _parse_docx(path: Path) -> ParsedFile:
    """Parse a DOCX using ``python-docx``.

    Returns paragraphs as text and embedded tables as 3-D string lists.
    """
    import docx as python_docx

    document = python_docx.Document(str(path))

    paragraphs: list[str] = []
    for idx, para in enumerate(document.paragraphs):
        if idx >= _MAX_DOCX_PARAGRAPHS:
            paragraphs.append("... (段落已截断)")
            break
        text = (para.text or "").strip()
        if text:
            paragraphs.append(text)

    tables: list[list[list[str]]] = []
    for table in document.tables:
        rendered = [
            [(cell.text or "").strip() for cell in row.cells] for row in table.rows
        ]
        if any(any(c for c in row) for row in rendered):
            tables.append(rendered)

    return ParsedFile(
        kind="docx",
        pages=len(paragraphs),
        text="\n".join(paragraphs),
        tables=tables,
    )


def _parse_xlsx(path: Path) -> ParsedFile:
    """Parse an XLSX using ``openpyxl`` in read-only mode.

    Every sheet becomes one entry in :attr:`ParsedFile.tables`, prefixed
    by a header row containing the sheet name so downstream prompts can
    differentiate.
    """
    from openpyxl import load_workbook

    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)

    tables: list[list[list[str]]] = []
    text_lines: list[str] = []
    for sheet in workbook.worksheets:
        sheet_rows: list[list[str]] = [[f"[Sheet] {sheet.title}"]]
        text_lines.append(f"[Sheet] {sheet.title}")
        for ridx, row in enumerate(sheet.iter_rows(values_only=True)):
            if ridx >= _MAX_XLSX_ROWS_PER_SHEET:
                sheet_rows.append(["... (rows truncated)"])
                text_lines.append("... (rows truncated)")
                break
            cells = ["" if c is None else str(c).strip() for c in row]
            if any(cells):
                sheet_rows.append(cells)
                text_lines.append(" | ".join(cells))
        tables.append(sheet_rows)

    return ParsedFile(
        kind="xlsx",
        pages=len(workbook.sheetnames),
        text="\n".join(text_lines),
        tables=tables,
    )


async def parse_file(path: Path) -> ParsedFile:
    """Dispatch on file extension and return a :class:`ParsedFile`.

    Args:
        path: Local filesystem path to the uploaded document.

    Returns:
        A populated :class:`ParsedFile`.

    Raises:
        ValueError: If the extension is not one of ``.pdf / .docx / .xlsx``.
        FileNotFoundError: If ``path`` does not exist.
    """
    if not path.exists():
        raise FileNotFoundError(f"file does not exist: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".xlsx":
        return _parse_xlsx(path)
    raise ValueError(f"unsupported file format: {suffix}")


# --------------------------------------------------------------------------- #
# Summariser                                                                  #
# --------------------------------------------------------------------------- #


def _render_tables_for_prompt(tables: list[list[list[str]]], cap: int = 3) -> str:
    """Render the first few tables as pipe-delimited text for the prompt."""
    if not tables:
        return ""
    rendered: list[str] = []
    for tidx, table in enumerate(tables[:cap]):
        rendered.append(f"[Table {tidx + 1}]")
        for row in table[:30]:
            rendered.append(" | ".join(row))
    return "\n".join(rendered)


async def _summarize(parsed: ParsedFile) -> str:
    """Ask a light model to summarise the parsed document.

    The prompt biases the model toward numbers, tabular data, and any
    environment / emissions context — that's what users typically upload
    in this product.
    """
    text_blob = parsed.text[:_MAX_TEXT_CHARS]
    tables_blob = _render_tables_for_prompt(parsed.tables)

    payload_sections: list[str] = []
    if text_blob:
        payload_sections.append(f"## 正文\n{text_blob}")
    if tables_blob:
        payload_sections.append(f"## 表格\n{tables_blob}")
    payload = "\n\n".join(payload_sections) or "(空文档)"

    system_prompt = (
        "你是一个化工碳数据领域的文档摘要助手。请基于给定文档内容生成一份"
        "不超过 300 字的中文摘要，重点关注：1) 数据表中的指标与数值；"
        "2) 涉及碳排放、能耗、污染物、履约的关键事实；3) 数据所属企业与年份。"
        "若文档与碳/能耗无关，简述其主题与关键信息即可。不要编造未出现的数字。"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": payload},
    ]

    model = get_light_model(temperature=0.2, max_tokens=600)
    try:
        response = await safe_ainvoke(model, messages)
        raw = response.content if isinstance(response.content, str) else str(response.content)
        summary = strip_thinking(raw).strip()
    except Exception as e:  # noqa: BLE001
        _log.warning("file summariser failed: %s", e)
        return (
            "文件已成功解析，但摘要生成失败。"
            f"识别到 {parsed.pages} 页/段，"
            f"{len(parsed.tables)} 张表格。请告诉我您想了解的具体内容。"
        )

    if not summary:
        return f"文件已解析，共 {parsed.pages} 页/段，{len(parsed.tables)} 张表格。"
    return summary


# --------------------------------------------------------------------------- #
# Node                                                                        #
# --------------------------------------------------------------------------- #


def _resolve_file_path(state: AgentState) -> str:
    """Return the uploaded file path from state, preferring explicit keys."""
    explicit = cast(Any, state).get("uploaded_file_path")
    if explicit:
        return str(explicit)
    cfg = state.get("config") or {}
    return str(cfg.get("file_path") or "")


async def file_query_node(
    state: AgentState,
    config: RunnableConfig,
) -> dict[str, Any]:
    """Parse an uploaded file and emit a summary to the user.

    On parsing failures the node returns a polite error so the graph never
    crashes mid-conversation.
    """
    del config

    file_path = _resolve_file_path(state)
    if not file_path:
        return {
            "messages": [AIMessage(content="未检测到上传文件路径，请先上传文件后再提问。")],
            "agent_results": {
                "file_agent": {
                    "agent": "file_agent",
                    "status": "error",
                    "data": {},
                    "message": "missing file path",
                }
            },
        }

    path = Path(file_path)
    if not path.exists():
        return {
            "messages": [AIMessage(content="文件不存在或已失效，请重新上传后再试。")],
            "agent_results": {
                "file_agent": {
                    "agent": "file_agent",
                    "status": "error",
                    "data": {"file_path": file_path},
                    "message": "file not found",
                }
            },
        }

    try:
        parsed = await parse_file(path)
    except ValueError as e:
        _log.info("file_query_node unsupported format: %s", e)
        return {
            "messages": [AIMessage(content=_FALLBACK_REPLY)],
            "agent_results": {
                "file_agent": {
                    "agent": "file_agent",
                    "status": "error",
                    "data": {"file_path": file_path, "suffix": path.suffix.lower()},
                    "message": str(e),
                }
            },
        }
    except Exception as e:  # noqa: BLE001
        _log.exception("file_query_node parse error: %s", e)
        return {
            "messages": [
                AIMessage(content="文件解析过程中出现问题，请确认文件未损坏后重试。")
            ],
            "agent_results": {
                "file_agent": {
                    "agent": "file_agent",
                    "status": "error",
                    "data": {"file_path": file_path},
                    "message": f"parse_error:{type(e).__name__}",
                }
            },
        }

    summary = await _summarize(parsed)

    return {
        "messages": [AIMessage(content=summary)],
        "agent_results": {
            "file_agent": {
                "agent": "file_agent",
                "status": "success",
                "data": {
                    "file_path": file_path,
                    "kind": parsed.kind,
                    "pages": parsed.pages,
                    "tables": len(parsed.tables),
                },
                "message": "file parsed and summarised",
            }
        },
    }
